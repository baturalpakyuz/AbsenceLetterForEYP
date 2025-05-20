import os
import shutil
import cloudconvert
from docx import Document
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QLabel, 
                             QLineEdit, QPushButton, QTableWidget, QTableWidgetItem, QFileDialog,
                             QProgressBar, QMessageBox, QHeaderView)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from pathvalidate import sanitize_filename


class DocumentWorker(QThread):
    progress = pyqtSignal(int)
    message = pyqtSignal(str)
    finished = pyqtSignal()
    error = pyqtSignal(str)

    def __init__(self, config):
        super().__init__()
        self.config = config
        self.running = True

    def run(self):
        try:
            cloudconvert.configure(api_key=self.config['api_key'])
            total = len(self.config['participants'])

            for idx, (name, is_delegate) in enumerate(self.config['participants']):
                if not self.running:
                    break

                # Sanitize filename
                safe_name = sanitize_filename(name, replacement_text='_')
                filename = f"{safe_name}_AbsenceLetter.docx"
                docx_path = os.path.join(self.config['output_folder'], filename)

                # Configure replacements
                replacements = {
                    'xxxxx': name,
                    'ttttt': self.config['conference_name'],
                    'ddddd': self.config['delegate_dates'] if is_delegate else self.config['official_dates']
                }

                # Generate DOCX
                try:
                    shutil.copy2(self.config['template_path'], docx_path)
                    doc = Document(docx_path)

                    # Replace placeholders
                    for paragraph in doc.paragraphs:
                        for placeholder, value in replacements.items():
                            self.replace_text_in_runs(paragraph, placeholder, value)

                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    for placeholder, value in replacements.items():
                                        self.replace_text_in_runs(paragraph, placeholder, value)

                    doc.save(docx_path)
                    self.message.emit(f"Created DOCX: {docx_path}")

                    # Convert to PDF
                    if self.convert_to_pdf(docx_path):
                        pdf_path = docx_path.replace('.docx', '.pdf')
                        self.message.emit(f"Converted to PDF: {pdf_path}")

                    self.progress.emit(int((idx + 1) / total * 100))

                except Exception as e:
                    self.error.emit(f"Error processing {name}: {e}")

            self.finished.emit()

        except Exception as e:
            self.error.emit(f"Critical error: {e}")

    def replace_text_in_runs(self, paragraph, old_text, new_text):
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)

    def convert_to_pdf(self, docx_path):
        try:
            pdf_path = docx_path.replace('.docx', '.pdf')
            job = cloudconvert.Job.create(payload={
                "tasks": {
                    'upload': {'operation': 'import/upload'},
                    'convert': {
                        'operation': 'convert',
                        'input': 'upload',
                        'output_format': 'pdf',
                        'engine': 'office'
                    },
                    'export': {
                        'operation': 'export/url',
                        'input': 'convert',
                        'inline': False,
                        'archive_multiple_files': False
                    }
                }
            })

            upload_task_id = job['tasks'][0]['id']
            upload_task = cloudconvert.Task.find(id=upload_task_id)
            cloudconvert.Task.upload(file_name=docx_path, task=upload_task)
            job = cloudconvert.Job.wait(id=job['id'])

            for task in job['tasks']:
                if task.get('operation') == 'export/url' and task['status'] == 'finished':
                    cloudconvert.download(filename=pdf_path, url=task['result']['files'][0]['url'])
                    return True
            return False

        except Exception as e:
            self.error.emit(f"Conversion failed: {e}")
            return False

    def stop(self):
        self.running = False


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Document Generator")
        self.setGeometry(100, 100, 800, 600)
        self.worker = None
        self.output_folder = os.path.join(os.path.expanduser("~"), "DocumentGeneratorOutput")

        # Create main widget and layout
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        layout = QVBoxLayout(main_widget)

        # API Key Section
        api_layout = QHBoxLayout()
        api_layout.addWidget(QLabel("CloudConvert API Key:"))
        self.api_key_input = QLineEdit()
        self.api_key_input.setPlaceholderText("Enter your API key here...")
        self.api_key_input.setEchoMode(QLineEdit.Password)
        api_layout.addWidget(self.api_key_input)
        self.save_api_btn = QPushButton("Save Key")
        self.save_api_btn.clicked.connect(self.save_api_key)
        api_layout.addWidget(self.save_api_btn)
        layout.addLayout(api_layout)

        # Template Section
        template_layout = QHBoxLayout()
        self.template_btn = QPushButton("Select Template")
        self.template_btn.clicked.connect(self.select_template)
        template_layout.addWidget(self.template_btn)
        self.template_label = QLabel("No template selected")
        template_layout.addWidget(self.template_label)
        layout.addLayout(template_layout)

        # Output Directory Section
        outdir_layout = QHBoxLayout()
        self.outdir_btn = QPushButton("Select Output Directory")
        self.outdir_btn.clicked.connect(self.select_output_folder)
        outdir_layout.addWidget(self.outdir_btn)
        self.outdir_label = QLabel(self.output_folder)
        outdir_layout.addWidget(self.outdir_label)
        layout.addLayout(outdir_layout)

        # Conference Details
        form_layout = QVBoxLayout()
        self.conference_name = QLineEdit()
        self.official_dates = QLineEdit()
        self.delegate_dates = QLineEdit()

        form_layout.addWidget(QLabel("Conference Name:"))
        form_layout.addWidget(self.conference_name)
        form_layout.addWidget(QLabel("Official Dates (DD/MM/YYYY-DD/MM/YYYY):"))
        form_layout.addWidget(self.official_dates)
        form_layout.addWidget(QLabel("Delegate Dates (DD/MM/YYYY-DD/MM/YYYY):"))
        form_layout.addWidget(self.delegate_dates)
        layout.addLayout(form_layout)

        # Participants Table
        self.table = QTableWidget(0, 2)
        self.table.setHorizontalHeaderLabels(["Name", "Delegate"])
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        layout.addWidget(QLabel("Participants:"))
        layout.addWidget(self.table)

        # Table Controls
        table_btn_layout = QHBoxLayout()
        self.add_btn = QPushButton("Add Participant")
        self.add_btn.clicked.connect(self.add_participant)
        self.remove_btn = QPushButton("Remove Selected")
        self.remove_btn.clicked.connect(self.remove_participant)
        table_btn_layout.addWidget(self.add_btn)
        table_btn_layout.addWidget(self.remove_btn)
        layout.addLayout(table_btn_layout)

        # Progress
        self.progress = QProgressBar()
        layout.addWidget(self.progress)

        # Log/Status
        self.status_label = QLabel()
        layout.addWidget(self.status_label)

        # Generate Button
        self.generate_btn = QPushButton("Generate Documents")
        self.generate_btn.clicked.connect(self.start_generation)
        layout.addWidget(self.generate_btn)

        self.load_settings()

    def select_output_folder(self):
        path = QFileDialog.getExistingDirectory(self, "Select Output Directory", self.output_folder)
        if path:
            self.output_folder = path
            self.outdir_label.setText(path)

    def add_participant(self):
        row = self.table.rowCount()
        self.table.insertRow(row)

        name_item = QTableWidgetItem()
        delegate_item = QTableWidgetItem()
        delegate_item.setFlags(Qt.ItemIsUserCheckable | Qt.ItemIsEnabled)
        delegate_item.setCheckState(Qt.Unchecked)

        self.table.setItem(row, 0, name_item)
        self.table.setItem(row, 1, delegate_item)

    def remove_participant(self):
        selected = self.table.selectedItems()
        if selected:
            self.table.removeRow(selected[0].row())

    def select_template(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select Template", "", "Word Documents (*.docx)")
        if path:
            self.template_path = path
            self.template_label.setText(os.path.basename(path))

    def save_api_key(self):
        # In a real app, you should encrypt this
        with open("config.ini", "w") as f:
            f.write(self.api_key_input.text())
        QMessageBox.information(self, "Success", "API key saved!")

    def load_settings(self):
        try:
            with open("config.ini", "r") as f:
                self.api_key_input.setText(f.read())
        except FileNotFoundError:
            pass

    def validate_inputs(self):
        if not self.api_key_input.text():
            QMessageBox.critical(self, "Error", "API key is required!")
            return False
        if not hasattr(self, 'template_path') or not self.template_path.endswith(".docx"):
            QMessageBox.critical(self, "Error", "Valid template file is required!")
            return False
        if not self.conference_name.text():
            QMessageBox.critical(self, "Error", "Conference name is required!")
            return False
        if self.table.rowCount() == 0:
            QMessageBox.critical(self, "Error", "Add at least one participant!")
            return False
        return True

    def start_generation(self):
        if not self.validate_inputs():
            return

        config = {
            'api_key': self.api_key_input.text(),
            'template_path': self.template_path,
            'conference_name': self.conference_name.text(),
            'official_dates': self.official_dates.text(),
            'delegate_dates': self.delegate_dates.text(),
            'output_folder': self.output_folder,
            'participants': []
        }

        # Create output directory
        os.makedirs(config['output_folder'], exist_ok=True)

        # Get participants from table
        for row in range(self.table.rowCount()):
            name = self.table.item(row, 0).text()
            is_delegate = self.table.item(row, 1).checkState() == Qt.Checked
            config['participants'].append((name, is_delegate))

        self.worker = DocumentWorker(config)
        self.worker.progress.connect(self.progress.setValue)
        self.worker.message.connect(self.status_label.setText)
        self.worker.finished.connect(self.on_finished)
        self.worker.error.connect(self.show_error)

        self.generate_btn.setEnabled(False)
        self.worker.start()

    def show_error(self, message):
        QMessageBox.critical(self, "Error", message)
        self.on_finished()

    def on_finished(self):
        self.generate_btn.setEnabled(True)
        self.progress.setValue(0)
        QMessageBox.information(self, "Complete", "Document generation finished!")

    def closeEvent(self, event):
        if self.worker and self.worker.isRunning():
            self.worker.stop()
            self.worker.quit()
            self.worker.wait()
        event.accept()


if __name__ == "__main__":
    app = QApplication([])
    window = MainWindow()
    window.show()
    app.exec_()